"""
Vedic RAG Demo - Retrieval-Augmented Generation for Sanskrit documents.
Modular pipeline: Document Loader -> Preprocessing & Chunking -> Embeddings -> FAISS -> Retriever -> flan-t5 Generator.
CPU-only; supports .txt, .pdf, .docx ingestion; Sanskrit / transliterated queries.
Uploaded files (PDF, TXT, DOCX) can be used as the query context on the website.
"""

from __future__ import annotations

import io
import os
import re
import time
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union

# ---------------------------------------------------------------------------
# Logging (performance and diagnostics)
# ---------------------------------------------------------------------------
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# 1. Document Loader - .txt and .pdf support
# ---------------------------------------------------------------------------

class DocumentLoader:
    """Load Sanskrit/English documents from .txt, .pdf, and .docx files."""

    @staticmethod
    def load_docx(path_or_file: Union[str, io.BytesIO]) -> List[Tuple[str, Dict[str, Any]]]:
        """Load a .docx file. path_or_file can be file path (str) or file-like (BytesIO)."""
        out = []
        try:
            from docx import Document
            if isinstance(path_or_file, (str, Path)):
                doc = Document(path_or_file)
                source = path_or_file
            else:
                doc = Document(path_or_file)
                source = getattr(path_or_file, "name", "uploaded.docx")
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            if paras:
                text = "\n\n".join(paras)
                out.append((text, {"source": str(source), "type": "docx"}))
        except ImportError:
            logger.warning("python-docx not installed. pip install python-docx for .docx support.")
        except Exception as e:
            logger.warning("Failed to load DOCX %s: %s", path_or_file, e)
        return out

    @staticmethod
    def load_txt(path: str) -> List[Tuple[str, Dict[str, Any]]]:
        """Load a single .txt file. Returns list of (text, metadata)."""
        out = []
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            if text.strip():
                out.append((text, {"source": path, "type": "txt"}))
        except Exception as e:
            logger.warning("Failed to load %s: %s", path, e)
        return out

    @staticmethod
    def load_pdf(path: str) -> List[Tuple[str, Dict[str, Any]]]:
        """Load a single .pdf file using pdfplumber. Returns list of (text, metadata)."""
        out = []
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages):
                    t = page.extract_text()
                    if t and t.strip():
                        out.append((t, {"source": path, "page": i + 1, "type": "pdf"}))
        except ImportError:
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(path)
                for i, page in enumerate(reader.pages):
                    t = page.extract_text()
                    if t and t.strip():
                        out.append((t, {"source": path, "page": i + 1, "type": "pdf"}))
            except Exception as e:
                logger.warning("PDF load failed for %s: %s", path, e)
        except Exception as e:
            logger.warning("Failed to load PDF %s: %s", path, e)
        return out

    @classmethod
    def load_directory(cls, directory: str) -> List[Tuple[str, Dict[str, Any]]]:
        """Load all .txt, .pdf, and .docx files from a directory. Returns list of (text, metadata)."""
        directory = os.path.abspath(directory)
        if not os.path.isdir(directory):
            return []
        documents = []
        for fname in sorted(os.listdir(directory)):
            path = os.path.join(directory, fname)
            if not os.path.isfile(path):
                continue
            low = fname.lower()
            if low.endswith(".txt"):
                documents.extend(cls.load_txt(path))
            elif low.endswith(".pdf"):
                documents.extend(cls.load_pdf(path))
            elif low.endswith(".docx"):
                documents.extend(cls.load_docx(path))
        return documents

    @classmethod
    def load_uploaded_files(
        cls, files: List[Any]
    ) -> List[Tuple[str, Dict[str, Any]]]:
        """
        Load from Streamlit UploadedFile or list of (name, bytes).
        Supports .txt, .pdf, .docx. Returns list of (text, metadata).
        """
        documents = []
        for f in files:
            if hasattr(f, "name") and hasattr(f, "read"):
                name = getattr(f, "name", "upload")
                f.seek(0)
                raw = f.read()
                if isinstance(raw, str):
                    raw = raw.encode("utf-8", errors="replace")
                bio = io.BytesIO(raw)
            elif isinstance(f, (list, tuple)) and len(f) >= 2:
                name, raw = f[0], f[1]
                if isinstance(raw, str):
                    raw = raw.encode("utf-8", errors="replace")
                bio = io.BytesIO(raw)
            else:
                continue
            low = name.lower()
            if low.endswith(".txt"):
                try:
                    text = bio.read().decode("utf-8", errors="replace")
                    if text.strip():
                        documents.append((text, {"source": name, "type": "txt"}))
                except Exception as e:
                    logger.warning("Failed to read TXT %s: %s", name, e)
            elif low.endswith(".pdf"):
                try:
                    import tempfile
                    bio.seek(0)
                    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                        tmp.write(bio.read())
                        tmp_path = tmp.name
                    try:
                        for text, meta in cls.load_pdf(tmp_path):
                            documents.append((text, {**meta, "source": name}))
                    finally:
                        try:
                            os.unlink(tmp_path)
                        except Exception:
                            pass
                except Exception as e:
                    logger.warning("Failed to read PDF %s: %s", name, e)
            elif low.endswith(".docx"):
                bio.seek(0)
                for text, meta in cls.load_docx(bio):
                    documents.append((text, {**meta, "source": name}))
        return documents


# ---------------------------------------------------------------------------
# 2. Preprocessing - clean Sanskrit text, chunk by tokens (500 / 50 overlap)
# ---------------------------------------------------------------------------

class Preprocessor:
    """Clean Sanskrit text and chunk with token overlap."""

    CHUNK_SIZE = 500
    CHUNK_OVERLAP = 50

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self._tokenizer = None

    def _get_tokenizer(self):
        """Lazy load tokenizer for token-based chunking (same as embedding model where possible)."""
        if self._tokenizer is None:
            try:
                from transformers import AutoTokenizer
                self._tokenizer = AutoTokenizer.from_pretrained(
                    "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
                )
            except Exception:
                self._tokenizer = None
        return self._tokenizer

    @staticmethod
    def clean_sanskrit(text: str) -> str:
        """Clean Sanskrit text: normalize spaces, remove control chars, strip."""
        if not text or not isinstance(text, str):
            return ""
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    def chunk_text(self, text: str, metadata: Optional[Dict] = None) -> List[Dict[str, Any]]:
        """Split text into chunks of chunk_size tokens with chunk_overlap overlap.
        Uses word-based chunking (no tokenizer load) for fast index build; ~500 words ≈ 500 tokens."""
        text = self.clean_sanskrit(text)
        if not text:
            return []
        meta = metadata or {}
        words = text.split()
        chunks = []
        start = 0
        while start < len(words):
            end = min(start + self.chunk_size, len(words))
            chunk_text = self.clean_sanskrit(" ".join(words[start:end]))
            if chunk_text:
                chunks.append({"text": chunk_text, "metadata": {**meta}})
            if end >= len(words):
                break
            start = max(end - self.chunk_overlap, start + 1)
        return chunks

    def process_documents(
        self, documents: List[Tuple[str, Dict[str, Any]]]
    ) -> List[Dict[str, Any]]:
        """Turn (text, metadata) pairs into chunk list."""
        all_chunks = []
        for text, meta in documents:
            all_chunks.extend(self.chunk_text(text, meta))
        return all_chunks


# ---------------------------------------------------------------------------
# 3. Embeddings + 4. Vector Index (FAISS)
# ---------------------------------------------------------------------------

class EmbeddingIndex:
    """Multilingual embeddings (paraphrase-multilingual-MiniLM-L12-v2) and FAISS index. CPU only."""

    MODEL_NAME = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"

    def __init__(self, device: str = "cpu"):
        self.device = device
        self._model = None
        self._index = None
        self._chunks: List[Dict[str, Any]] = []
        self._index_path: Optional[str] = None

    def _get_model(self):
        if self._model is None:
            # prefer sentence-transformers if available, otherwise fall back to
            # huggingface transformers with manual mean-pooling.
            try:
                from sentence_transformers import SentenceTransformer
                self._model = SentenceTransformer(self.MODEL_NAME, device=self.device)
                self._use_sentence_transformers = True
            except ImportError:
                from transformers import AutoModel, AutoTokenizer
                self._tokenizer = AutoTokenizer.from_pretrained(self.MODEL_NAME)
                self._model = AutoModel.from_pretrained(self.MODEL_NAME).to(self.device)
                self._use_sentence_transformers = False
        return self._model

    def build(self, chunks: List[Dict[str, Any]], index_dir: Optional[str] = None) -> None:
        """Build embeddings and FAISS index from chunks."""
        import numpy as np
        import faiss
        texts = [c["text"] for c in chunks]
        model = self._get_model()
        logger.info("Computing embeddings for %d chunks (CPU)...", len(texts))
        if getattr(self, '_use_sentence_transformers', False):
            vecs = model.encode(texts, show_progress_bar=len(texts) > 50, device=self.device)
        else:
            # manual encode via transformers
            import torch
            encoded = self._tokenizer(texts, padding=True, truncation=True, return_tensors='pt')
            encoded = {k: v.to(self.device) for k, v in encoded.items()}
            with torch.no_grad():
                out = model(**encoded)
            # mean pooling over tokens
            vecs = out.last_hidden_state.mean(dim=1).cpu().numpy()
        vecs = np.array(vecs, dtype=np.float32)
        d = vecs.shape[1]
        index = faiss.IndexFlatL2(d)
        index.add(vecs)
        self._index = index
        self._chunks = chunks
        if index_dir:
            os.makedirs(index_dir, exist_ok=True)
            faiss.write_index(index, os.path.join(index_dir, "faiss.index"))
            import json
            meta_path = os.path.join(index_dir, "chunks_metadata.json")
            with open(meta_path, "w", encoding="utf-8") as f:
                json.dump([{"text": c["text"], "metadata": c.get("metadata", {})} for c in chunks], f, ensure_ascii=False)
            self._index_path = index_dir
        logger.info("FAISS index built: %d vectors", len(chunks))

    def load(self, index_dir: str) -> bool:
        """Load existing FAISS index and chunk metadata."""
        import numpy as np
        import faiss
        idx_path = os.path.join(index_dir, "faiss.index")
        meta_path = os.path.join(index_dir, "chunks_metadata.json")
        if not os.path.isfile(idx_path) or not os.path.isfile(meta_path):
            return False
        self._index = faiss.read_index(idx_path)
        import json
        with open(meta_path, "r", encoding="utf-8") as f:
            self._chunks = json.load(f)
        self._index_path = index_dir
        logger.info("Loaded FAISS index from %s (%d chunks)", index_dir, len(self._chunks))
        return True

    def search(self, query: str, k: int = 3) -> List[Dict[str, Any]]:
        """Return top-k chunks for query. Requires build() or load() already done."""
        if self._index is None or not self._chunks:
            return []
        model = self._get_model()
        if getattr(self, '_use_sentence_transformers', False):
            q_vec = model.encode([query], device=self.device)
        else:
            import torch
            encoded = self._tokenizer([query], padding=True, truncation=True, return_tensors='pt')
            encoded = {k: v.to(self.device) for k, v in encoded.items()}
            with torch.no_grad():
                out = model(**encoded)
            q_vec = out.last_hidden_state.mean(dim=1).cpu().numpy()
        import numpy as np
        q_vec = np.array(q_vec, dtype=np.float32)
        scores, indices = self._index.search(q_vec, min(k, len(self._chunks)))
        out = []
        for i, idx in enumerate(indices[0]):
            if 0 <= idx < len(self._chunks):
                ch = self._chunks[idx]
                c = dict(ch) if isinstance(ch, dict) else {"text": getattr(ch, "text", ""), "metadata": getattr(ch, "metadata", {})}
                c["score"] = float(scores[0][i])
                out.append(c)
        return out


# ---------------------------------------------------------------------------
# 5. Retriever - top-k (k=3)
# ---------------------------------------------------------------------------

class Retriever:
    """Retrieve top-k relevant chunks from FAISS index."""

    def __init__(self, index: EmbeddingIndex, top_k: int = 3):
        self.index = index
        self.top_k = top_k

    def retrieve(self, query: str) -> List[Dict[str, Any]]:
        return self.index.search(query, k=self.top_k)


# ---------------------------------------------------------------------------
# 6. Generator - flan-t5-base, CPU
# ---------------------------------------------------------------------------

class Generator:
    """Answer generation using google/flan-t5-base. CPU only."""

    MODEL_NAME = "google/flan-t5-base"

    def __init__(self, device: str = "cpu"):
        self.device = device
        self._model = None
        self._tokenizer = None

    def _get_model_and_tokenizer(self):
        if self._model is None:
            from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
            self._tokenizer = AutoTokenizer.from_pretrained(self.MODEL_NAME)
            self._model = AutoModelForSeq2SeqLM.from_pretrained(self.MODEL_NAME)
            self._model = self._model.to(self.device)
            self._model.eval()
        return self._model, self._tokenizer

    def generate(self, context: str, question: str, max_new_tokens: int = 150) -> str:
        """Generate answer given context and question. Uses assignment prompt format."""
        prompt = f"""Context:\n{context}\n\nQuestion:\n{question}\n\nAnswer in Sanskrit or English explanation."""
        model, tokenizer = self._get_model_and_tokenizer()
        inputs = tokenizer(prompt, return_tensors="pt", truncation=True, max_length=512).input_ids.to(self.device)
        out = model.generate(inputs, max_new_tokens=max_new_tokens, do_sample=False, num_beams=2)
        return tokenizer.decode(out[0], skip_special_tokens=True)


# ---------------------------------------------------------------------------
# Orchestrator: VedicRAGDemo (full pipeline + performance logging)
# ---------------------------------------------------------------------------

def _find_corpus_dir() -> Optional[str]:
    """Resolve corpus directory (nalanda_library or vedic_texts). Returns None if not found."""
    for name in ("nalanda_library", "vedic_texts"):
        if os.path.exists(name):
            return os.path.abspath(name)
        parent = os.path.join(os.path.dirname(os.getcwd()), name)
        if os.path.exists(parent):
            return os.path.abspath(parent)
    return None


class VedicRAGDemo:
    """
    Full RAG pipeline: Loader -> Preprocessing -> Embedding -> FAISS -> Retriever -> flan-t5.
    CPU-only. Supports Sanskrit and transliterated queries.
    """

    def __init__(
        self,
        corpus_dir: Optional[str] = None,
        index_dir: Optional[str] = None,
        device: str = "cpu",
        top_k: int = 3,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
    ):
        self.device = device
        self.corpus_dir = corpus_dir if corpus_dir is not None else _find_corpus_dir()
        self.index_dir = (index_dir or (os.path.join(self.corpus_dir, "faiss_index") if self.corpus_dir else None))
        self.top_k = top_k
        self.loader = DocumentLoader()
        self.preprocessor = Preprocessor(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        self.embedding_index = EmbeddingIndex(device=device)
        self.retriever = Retriever(self.embedding_index, top_k=top_k)
        self.generator = Generator(device=device)
        self._index_ready = False
        self._upload_embedding_index: Optional[EmbeddingIndex] = None

    def _get_effective_retriever(self) -> Retriever:
        """Use upload index if available, else main corpus index."""
        if self._upload_embedding_index is not None and getattr(self._upload_embedding_index, "_chunks", None):
            return Retriever(self._upload_embedding_index, top_k=self.top_k)
        return self.retriever

    def build_from_uploaded_files(self, files: List[Any]) -> None:
        """
        Build an in-memory FAISS index from uploaded files (PDF, TXT, DOCX).
        After this, queries use the uploaded documents until clear_uploaded_documents() or new uploads.
        """
        documents = self.loader.load_uploaded_files(files)
        if not documents:
            raise ValueError("No valid content from uploaded files. Support: .txt, .pdf, .docx")
        chunks = self.preprocessor.process_documents(documents)
        if not chunks:
            raise ValueError("No chunks from uploaded documents.")
        self._upload_embedding_index = EmbeddingIndex(device=self.device)
        self._upload_embedding_index.build(chunks, index_dir=None)
        logger.info("Built upload index from %d chunks.", len(chunks))

    def clear_uploaded_documents(self) -> None:
        """Switch back to default corpus (if any)."""
        self._upload_embedding_index = None

    def has_uploaded_documents(self) -> bool:
        """True if currently using uploaded documents for retrieval."""
        return self._upload_embedding_index is not None and bool(getattr(self._upload_embedding_index, "_chunks", None))

    def _ensure_index(self) -> None:
        """Load existing FAISS index or build from corpus. Skipped when using uploaded documents."""
        if self.has_uploaded_documents():
            return
        if self._index_ready:
            return
        if self.corpus_dir is None or self.index_dir is None:
            raise FileNotFoundError(
                "Upload PDF/TXT/DOCX files above, or add documents to nalanda_library/ and run VedicDatasetGenerator.py."
            )
        if self.embedding_index.load(self.index_dir):
            self._index_ready = True
            return
        logger.info("Building FAISS index from corpus (this may take a few minutes on first run)...")
        documents = self.loader.load_directory(self.corpus_dir)
        if not documents:
            logger.warning("No .txt/.pdf/.docx in %s; checking for legacy verse .txt files.", self.corpus_dir)
            documents = self._load_legacy_corpus()
        if not documents:
            raise FileNotFoundError(
                f"No documents in {self.corpus_dir}. Add .txt/.pdf/.docx or run VedicDatasetGenerator.py."
            )
        chunks = self.preprocessor.process_documents(documents)
        if not chunks:
            raise ValueError("No chunks produced from documents.")
        self.embedding_index.build(chunks, self.index_dir)
        self._index_ready = True

    def _load_legacy_corpus(self) -> List[Tuple[str, Dict[str, Any]]]:
        """Load existing nalanda_library verse .txt files as single doc per file."""
        pattern = re.compile(
            r"Source:\s*(.+?)\nCategory:\s*(.+?)\nTitle:\s*(.+?)\n"
            r"Verse:\s*(.+?)\nSanskrit:\s*(.+?)\nEnglish:\s*(.+?)\n---",
            re.DOTALL,
        )
        documents = []
        for fname in sorted(os.listdir(self.corpus_dir)):
            if not fname.endswith(".txt"):
                continue
            path = os.path.join(self.corpus_dir, fname)
            if not os.path.isfile(path):
                continue
            try:
                with open(path, "r", encoding="utf-8") as f:
                    text = f.read()
                documents.append((text, {"source": path, "type": "txt"}))
            except Exception:
                continue
        return documents

    def query(self, user_query: str) -> Dict[str, Any]:
        """
        Run full RAG: retrieve top-k chunks, then generate answer.
        Returns dict with answer, sources, retrieval_time, generation_time, total_time.
        """
        self._ensure_index()
        t0 = time.perf_counter()
        retrieval_start = time.perf_counter()
        chunks = self._get_effective_retriever().retrieve(user_query)
        retrieval_time = time.perf_counter() - retrieval_start
        context = "\n\n".join([c.get("text", "") for c in chunks])
        if not context.strip():
            return {
                "answer": "No relevant context found for your query.",
                "sources": [],
                "retrieval_time": retrieval_time,
                "generation_time": 0.0,
                "total_time": time.perf_counter() - t0,
            }
        gen_start = time.perf_counter()
        answer = self.generator.generate(context, user_query)
        generation_time = time.perf_counter() - gen_start
        total_time = time.perf_counter() - t0
        logger.info(
            "query_latency=%.3fs retrieval=%.3fs generation=%.3fs",
            total_time, retrieval_time, generation_time,
        )
        return {
            "answer": answer,
            "sources": chunks,
            "retrieval_time": retrieval_time,
            "generation_time": generation_time,
            "total_time": total_time,
        }

    def retrieve_only(self, user_query: str) -> Tuple[List[Dict[str, Any]], float]:
        """Retrieve top-k chunks and return (sources, retrieval_time_seconds)."""
        self._ensure_index()
        t0 = time.perf_counter()
        chunks = self._get_effective_retriever().retrieve(user_query)
        elapsed = time.perf_counter() - t0
        return chunks, elapsed

    def search_verses(self, query: str, max_results: int = 5) -> List[Dict[str, Any]]:
        """
        Retriever-only: return top chunks formatted as verse-like dicts for UI compatibility.
        """
        self._ensure_index()
        chunks = self._get_effective_retriever().retrieve(query)
        out = []
        for c in chunks[:max_results]:
            text = c.get("text", "")
            meta = c.get("metadata", {})
            out.append({
                "title": meta.get("title", "Chunk"),
                "source": meta.get("source", "Document"),
                "category": meta.get("category", ""),
                "sanskrit": text[:200] + ("..." if len(text) > 200 else ""),
                "english": text,
            })
        return out

    def generate_response(self, query: str, max_verses: int = 3) -> str:
        """Full RAG response (answer + timing). For CLI."""
        result = self.query(query)
        lines = [result["answer"], ""]
        if result.get("sources"):
            lines.append("Retrieved context:")
            for i, s in enumerate(result["sources"][:max_verses], 1):
                lines.append(f"  {i}. {s.get('text', '')[:200]}...")
        lines.append("")
        lines.append(f"[retrieval: {result['retrieval_time']:.2f}s | generation: {result['generation_time']:.2f}s | total: {result['total_time']:.2f}s]")
        return "\n".join(lines)


if __name__ == "__main__":
    import sys
    try:
        rag = VedicRAGDemo()
        q = sys.argv[1] if len(sys.argv) > 1 else "duty"
        print(f"Query: {q}\n")
        print(rag.generate_response(q, max_verses=3))
    except FileNotFoundError as e:
        print(f"Error: {e}")
        print("Run VedicDatasetGenerator.py or add .txt/.pdf to nalanda_library/")
